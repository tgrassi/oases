from docx import Document


def do_table(data, fname, style='Light Shading Accent 1'):
  # extract shape from data

  # create word document
  doc = Document()

  for d in data:
    rows, cols = len(d), len(d[0])

    # add table
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    for i, row in enumerate(table.rows):
      for j, cell in enumerate(row.cells):
        cell.text = str(d[i][j])

    paragraph = doc.add_paragraph()
    paragraph.text = "\n"

  # save to file
  doc.save(fname)
  print("table saved in " +  fname)


if __name__ == "__main__":


  # data: list of columns, each column is a list, first element is title
  data1 = [["Name", "A", "B", "C", "Bzzzz"],
          ["Value 1", 0.1, 0.3, 123.2, 111.4],
          ["Value 2", 1, 555, 123, 39.33, 333.3]]

  data2 = [["Name", "E", "F", "XXXX"],
          ["Value X", 0.1222, 110.3, 1e3],
          ["Value Y", 0.1333, 7839.3, 1e3+1],
          ["Value Z", 8271, "asdsa", 0000.1, "aslksajò"]]


  do_table([data1, data2], "table.docx")

